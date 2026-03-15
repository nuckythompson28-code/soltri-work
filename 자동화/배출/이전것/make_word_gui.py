import os
import sys
import random
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import messagebox, ttk

# 실행 파일 경로를 고정하는 로직 추가
def get_base_path():
    if hasattr(sys, '_MEIPASS'):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))

def replace_text_with_font(doc, old_text, new_text):
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for p in all_paragraphs:
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    val = "" if pd.isna(new_text) else str(new_text)
                    run.text = run.text.replace(old_text, val)
                    run.font.name = '바탕'
                    r = run._element.rPr
                    r.rFonts.set(qn('w:eastAsia'), '바탕')

def get_weather_data(month):
    temp_map = {1: (-5, 2), 2: (-3, 5), 3: (2, 11), 4: (8, 18), 5: (14, 24), 6: (19, 27),
                7: (23, 30), 8: (24, 31), 9: (18, 26), 10: (11, 20), 11: (4, 12), 12: (-2, 4)}
    weather_list = ["맑음", "맑음", "맑음", "흐림", "구름조금"]
    min_t, max_t = temp_map[month]
    return random.choice(weather_list), str(random.randint(min_t, max_t))

def run_process(start_str, end_str):
    base_path = get_base_path()
    template_path = os.path.join(base_path, 'template.docx')
    csv_path = os.path.join(base_path, 'data.csv')
    output_base_dir = os.path.join(base_path, "운영기록부_생성결과")

    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {template_path}")
        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"데이터 파일을 찾을 수 없습니다: {csv_path}")

        try:
            df = pd.read_csv(csv_path, encoding='cp949')
        except:
            df = pd.read_csv(csv_path, encoding='utf-8')
        
        df['날짜'] = pd.to_datetime(df['날짜']).dt.strftime('%Y-%m-%d')
        df.set_index('날짜', inplace=True)

        current_date = datetime.strptime(start_str, "%Y-%m-%d")
        end_limit = datetime.strptime(end_str, "%Y-%m-%d")
        days_ko = ["월", "화", "수", "목", "금", "토", "일"]

        count = 0
        while current_date <= end_limit:
            date_key = current_date.strftime('%Y-%m-%d')
            if date_key not in df.index:
                break

            month_dir = os.path.join(output_base_dir, f"{current_date.month:02d}월")
            if not os.path.exists(month_dir): 
                os.makedirs(month_dir)

            doc = Document(template_path)
            wth, tmp = get_weather_data(current_date.month)
            is_weekend = current_date.weekday() >= 5
            
            row = df.loc[date_key]
            if isinstance(row, pd.DataFrame): row = row.iloc[0]
            vals = {f"VAL{i}": row[f"전력{i}"] for i in range(1, 9)}
            t17, t18 = ("", "") if is_weekend else ("07:00 ~ 12:00, 13:00~17:00", "07:00 ~ 12:00, 13:00~18:00")

            replace_text_with_font(doc, "YYYY", str(current_date.year))
            replace_text_with_font(doc, "MM", str(current_date.month))
            replace_text_with_font(doc, "DD", str(current_date.day))
            replace_text_with_font(doc, "WD", days_ko[current_date.weekday()])
            replace_text_with_font(doc, "WTH", wth)
            replace_text_with_font(doc, "TMP", tmp)
            replace_text_with_font(doc, "TIME17", t17)
            replace_text_with_font(doc, "TIME18", t18)
            for k, v in vals.items(): replace_text_with_font(doc, k, v)

            doc.save(os.path.join(month_dir, f"배출방지시설운영기록부({current_date.strftime('%y%m%d')}).docx"))
            current_date += timedelta(days=1)
            count += 1
        
        messagebox.showinfo("완료", f"성공! '{output_base_dir}' 폴더에 {count}개 파일이 생성되었습니다.")
    except Exception as e:
        messagebox.showerror("오류", f"에러 발생: {str(e)}")

# --- GUI 구성 ---
root = tk.Tk()
root.title("운영기록부 자동 생성기 (2026)")
root.geometry("350x250")

tk.Label(root, text="생성 기간을 입력하세요", font=("맑은 고딕", 12, "bold")).pack(pady=10)
tk.Label(root, text="시작일 (YYYY-MM-DD)").pack()
ent_start = tk.Entry(root, justify='center')
ent_start.insert(0, "2026-01-01")
ent_start.pack(pady=5)
tk.Label(root, text="종료일 (YYYY-MM-DD)").pack()
ent_end = tk.Entry(root, justify='center')
ent_end.insert(0, "2026-12-31")
ent_end.pack(pady=5)

btn = tk.Button(root, text="워드 파일 생성 시작", command=lambda: run_process(ent_start.get(), ent_end.get()), 
                bg="#2196F3", fg="white", font=("맑은 고딕", 10, "bold"), width=20)
btn.pack(pady=20)

root.mainloop()